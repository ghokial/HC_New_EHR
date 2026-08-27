from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def extract(path: Path) -> dict:
    doc = Document(path)
    blocks: list[dict] = []
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            blocks.append({
                "kind": "paragraph",
                "index": i,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            })
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append({"kind": "table", "index": ti, "rows": rows})
    return {
        "source": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "blocks": blocks,
    }


def main() -> None:
    out_dir = Path(sys.argv[1])
    out_dir.mkdir(parents=True, exist_ok=True)
    for raw in sys.argv[2:]:
        path = Path(raw)
        payload = extract(path)
        target = out_dir / f"{path.stem}.json"
        target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        print(target)


if __name__ == "__main__":
    main()
